"""
ingestion/word_parser.py
Handles: .docx, .doc
Extracts: paragraphs, tables, headings.
Requires: python-docx
"""

from __future__ import annotations
import io
import pandas as pd
from ingestion.base_parser import BaseParser, ParsedDocument
from core.logger import get_logger

logger = get_logger(__name__)


class WordParser(BaseParser):

    def can_parse(self, filename: str, mime_type: str = "") -> bool:
        return (filename.lower().endswith(".docx") or
                filename.lower().endswith(".doc") or
                "wordprocessingml" in mime_type or
                "msword" in mime_type)

    def parse(self, source, filename: str = "") -> ParsedDocument:
        doc = ParsedDocument(source_name=filename, source_type="word")
        try:
            from docx import Document
            if hasattr(source, "read"):
                word_doc = Document(io.BytesIO(source.read()))
            else:
                word_doc = Document(source)

            # Extract paragraphs grouped into chunks
            current_chunk = []
            for para in word_doc.paragraphs:
                text = para.text.strip()
                if not text:
                    if current_chunk:
                        doc.text_chunks.append(" ".join(current_chunk))
                        current_chunk = []
                else:
                    current_chunk.append(text)
            if current_chunk:
                doc.text_chunks.append(" ".join(current_chunk))

            # Extract tables
            for t_idx, table in enumerate(word_doc.tables):
                try:
                    rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
                    if len(rows) >= 2:
                        df = pd.DataFrame(rows[1:], columns=rows[0])
                        doc.dataframes.append(df)
                        doc.table_names.append(f"table_{t_idx+1}")
                except Exception as e:
                    doc.add_warning(f"Word table {t_idx} failed: {e}")

            logger.info(f"Word parsed: {len(doc.text_chunks)} chunks, "
                        f"{len(doc.dataframes)} tables")

        except ImportError:
            doc.add_warning("python-docx not installed. Run: pip install python-docx")
        except Exception as e:
            doc.add_warning(f"Word parse error: {e}")

        return doc
